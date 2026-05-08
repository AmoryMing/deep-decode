"""Build 图文 docx for Opus 4.7 deep decode."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
ART = os.path.dirname(HERE)
OUT = os.path.join(ART, '【AI笔记0417】Claude-Opus-4-7不是顶配.docx')

doc = Document()

# set default font to PingFang SC
style = doc.styles['Normal']
style.font.name = 'PingFang SC'
style.font.size = Pt(11)
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts')
    rPr.insert(0, rFonts)
rFonts.set(qn('w:eastAsia'), 'PingFang SC')
rFonts.set(qn('w:ascii'), 'PingFang SC')

def set_font(run, name='PingFang SC', size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rPr.append(rFonts)

def h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_font(r, size=20, bold=True, color='1F1A14')
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)

def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=15, bold=True, color='2383E2')
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)

def para(text, bold=False, italic=False, size=11, color='37352F'):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    r.font.italic = italic
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.7
    return p

def quote(text, cite=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_font(r, size=11, color='6B6052')
    r.font.italic = True
    if cite:
        r2 = p.add_run(f'\n—— {cite}')
        set_font(r2, size=10, color='9B8B72')

def img(path, width_inches=6.0, caption=None):
    doc.add_picture(os.path.join(ART, path), width=Inches(width_inches))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption)
        set_font(r, size=9, color='9B8B72')
        r.font.italic = True

# ============ COVER ============
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('AI 前沿拆解')
set_font(r, size=13, color='9B8B72')
p.paragraph_format.space_before = Pt(40)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Claude Opus 4.7 不是顶配')
set_font(r, size=26, bold=True, color='1F1A14')
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('顶配被 Anthropic 锁起来了')
set_font(r, size=20, bold=True, color='2383E2')
p.paragraph_format.space_after = Pt(16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('一次发布背后的 AI 发布新范式 · 2026-04-17')
set_font(r, size=11, color='787066')

img('material/pngs/00_cover.png', 6.2)

doc.add_page_break()

# ============ Section 1 ============
h1('一、开场: Anthropic 发了一个"刻意降配"的版本')
para('4 月 16 日，Anthropic 发布 Claude Opus 4.7。这个版本的 SWE-bench Verified 分数是 87.6%，超过了所有公开可用的模型。但它不是 Anthropic 手上最强的——最强的那个叫 Claude Mythos Preview，SWE-bench Verified 考了 93.9%，然后被锁进了一个 11 家科技巨头组成的"玻璃翼联盟"（Project Glasswing），普通开发者拿不到。')
para('这是第一次，有厂商主动把"最强模型"藏起来，只给"第二强"发广泛许可证。这件事比 Opus 4.7 本身更值得拆开看。')

para('先看原文怎么说的。Anthropic 在公告里有一段措辞非常罕见:')
quote('"Opus 4.7 is the first such model: its cyber capabilities are not as advanced as those of Mythos Preview (indeed, during its training we experimented with efforts to differentially reduce these capabilities)."\n（Opus 4.7 是第一个这样的模型: 它的网络安全能力不像 Mythos Preview 那么先进——事实上，在训练过程中，我们专门尝试了差异化降低这些能力。）', cite='Anthropic 官方公告')

para('请注意"differentially reduce"（差异化降低）这个词。这不是"新模型恰好能力更弱"，而是"我们主动在训练阶段把某些能力切掉"。这是一场训练侧的外科手术。')

para('把 Mythos Preview 和 Project Glasswing 联盟两件事串起来，Anthropic 正在开创一种新的模型发布范式: 最强的那个不公开，能公开的是被刻意阉割过的版本。Opus 4.7 不是"Anthropic 的新旗舰"——它是 Anthropic 愿意让你碰的最强模型。', bold=False)

img('material/pngs/01_tiering.png', 6.2, caption='图 1 · Anthropic 的四层模型分级发布结构')

# ============ Section 2 ============
h1('二、三个转向: 为什么"稍微强一点"就值得写篇公告')
para('把 Mythos 的影子放一边，单看 Opus 4.7 相对 4.6 的提升，其实并不以 benchmark 峰值为主线。真正值得从业者盯住的是三个非常具体的方向转变。')

h2('转向一: 从"能力天花板"到"token 效率"')
para('原文里埋了一句听起来轻描淡写但杀伤力巨大的话: Opus 4.7 在 xhigh 效果档位跑 100k token，性能已经等同于 Opus 4.6 在 max 档位跑 200k token。')
para('换句话说，同样的任务质量，token 消耗减半。定价没变（输入 $5/M，输出 $25/M），但用一半的 token 就能达到同样效果。Replit 的 CEO Michele Catasta 说得很直白: "same quality at lower cost"。', bold=False)
para('这个转向背后是行业阶段的切换。大模型的 price-performance 曲线已经从"能力陡峭上升期"进入"效率边际优化期"。Anthropic 说 Opus 4.7 在 Rakuten-SWE-Bench 上解决了 3 倍于 Opus 4.6 的生产任务——他们不说"benchmark 分数高了多少"，而说"通过的真实任务多了几倍"。衡量单位从 percentage point 换成了 task throughput。')

h2('转向二: 从"一次响应"到"数小时持续工作"')
para('Cognition CEO Scott Wu 对 Opus 4.7 的评价是: "It works coherently for hours, pushes through hard problems rather than giving up"——它能连贯工作好几个小时，遇到硬问题不放弃。')
para('Opus 4.7 给长任务问题发了两个新武器: Task Budget（任务预算）——你给 Claude 一个大致的 token 预算，模型会看到一个"剩余预算"的 running countdown，并据此决定优先级；Auto Mode（自主模式）——仅限 Claude Code 的 Max 用户，允许 Claude 在权限范围内自己做决策、不打断。')
para('Anthropic 在把 agent 产品化的颗粒度从"一轮对话"推到"一个任务"。他们不再把 Claude 卖成"你问我答的助手"，而是卖成"你交任务我做几小时的同事"。')

img('material/pngs/02_three_shifts.png', 6.2, caption='图 2 · Opus 4.7 的三个方向转变')

h2('转向三: 从"模型即终点"到"命令化工作流"')
para('Opus 4.7 带来的 /ultrareview 命令是这一转向最明显的信号。这不是模型升级——这是把一个特定工作流打包成了命令。')
para('/ultrareview 会对代码改动做一次结构化审查，从架构、安全、性能、可维护性几个维度找单次 review 容易漏掉的深层问题。CodeRabbit 在官方 testimonial 里说 Opus 4.7 是"the sharpest model we\'ve tested"，recall 提升超过 10%。')
para('Pro 和 Max 用户每月送 3 次免费使用。这是一个极具信号的商业动作: 模型的差异化红利在变薄，产品层的差异化红利在变厚。同一套模型能力，谁把它组织成更好用的工作流，谁就赢。')

doc.add_page_break()

# ============ Section 3 ============
h1('三、Benchmark 真相: 亮眼数字背后的层次')

para('Opus 4.7 公布的 benchmark 数字确实好看: SWE-bench Pro 64.3% (+10.9pt over 4.6), SWE-bench Verified 87.6% (+6.8pt), CursorBench 70% vs 58%, 视觉敏锐度 98.5% vs 54.5%, Rakuten-SWE-Bench 生产任务 3 倍。')
para('但如果你把 Mythos Preview 的分数也放进来对比，故事就不一样了。Mythos 在 SWE-bench Verified 上是 93.9%，比 Opus 4.7 高 6.3 个百分点。在 Anthropic 官方公布的基准表里，Mythos 在 agentic coding、网络安全漏洞复现、工具使用等绝大多数项目上都是最高分——但这一栏只出现在"展示用的对比表"里，不出现在"你能买到的产品线"里。')

img('material/pngs/03_benchmarks.png', 6.2, caption='图 3 · Mythos Preview 在对比表里领先但不对外发售')

para('这带来一个很有意思的变化: benchmark 榜单不再是单纯的能力展示，它成了一个战略博弈工具。Anthropic 把 Mythos 的分数放进对比表里不是为了宣传，而是为了说明"我们手上的还有更强的，但我们选择不放出来"——这是一种 capability signaling + restraint signaling 的组合拳。')

# ============ Section 4 ============
h1('四、差异化训练: 一个被轻描淡写的技术转折')

para('差异化训练意味着什么? 意味着一个模型可以在训练阶段被"手术式"地削弱某些能力，同时保留其他能力。这在过去是非常难做的——能力之间是耦合的，砍一刀通常会伤到相邻的肌肉。')
para('这件事的副作用已经显现。公告的安全评估部分有一行相当诚实的承认: Opus 4.7 在 controlled substance harm-reduction advice（管制物质减害建议）上比 Opus 4.6 稍差——更容易提供"过度详细的减害建议"。这说明差异化训练的精度还没到"只砍该砍的"。', bold=False)

img('material/pngs/04_differential.png', 6.2, caption='图 4 · 差异化训练的能力手术及其副作用')

para('如果差异化训练成熟，未来可能看到的模式是: 一个底座模型，训练出 N 个版本，每个版本针对不同行业场景削减不同的风险能力。企业版砍掉娱乐类，医疗版砍掉金融类，教育版砍掉安全类——模型定制化的颗粒度从 prompt 层下沉到训练层。')

# ============ Section 5 ============
h1('五、一个新概念: Model Tiering')

para('把 Mythos Preview + Project Glasswing + Opus 4.7 + Cyber Verification Program 串起来看，Anthropic 实际上开创了一个 Model Tiering（模型分级发布）的新范式。')
para('Tier 0 内部不发布 / Tier 1 联盟限制发布 (Mythos) / Tier 2 广泛商业发布 (Opus 4.7) / Tier 3 专业豁免 (Cyber Verification Program)。这个四层结构让 Anthropic 一次性解决了几个矛盾: 既能向投资人展示前沿能力，又能向监管者展示克制；既能保留商业化的广泛可及性，又能满足安全关键场景的专业需求。', bold=False)

# ============ Section 6 ============
h1('六、盲区: 这次发布没说的事')
para('1. Mythos 的能力上限没公开。"autonomously developed next-gen offensive cyber capability"具体能做到什么程度? 官方只做了定性描述，没给定量边界。')
para('2. "数小时连贯工作"没有时长上限。Cognition 说 hours，但具体多少小时? 这让 hours 听起来是一个营销词而不是工程指标。')
para('3. 价格没降。但如果 xhigh 100k 的性能等同 max 200k，用户实际付费会减半。这是一种"隐性降价"。')
para('4. 差异化训练的技术细节没公开。这可能是 Anthropic 接下来两年最核心的技术护城河。')

# ============ Section 7 ============
h1('七、对 AI 从业者意味着什么')
para('对 AI 产品经理: 重新设计你的 agentic 工作流的"任务单位"。过去把一个复杂任务拆成 20 轮对话；现在可以考虑交给 Claude 一次 2 小时的任务，用 Task Budget 控制上限。产品形态从"聊天界面"变成"任务看板"。')
para('对技术架构师: 如果你现在在 Opus 4.6 的 max 档位上跑生产流量，立即测 Opus 4.7 的 xhigh 档位。效果持平就直接切——成本减半。')
para('对 Claude Code 重度用户: /ultrareview 加入 PR workflow，关键模块强制跑一次。Auto Mode 是 Max 专享，升级 ROI 比以往都高。')
para('对 AI 安全研究者: 关注 Cyber Verification Program 的准入机制。它是 Anthropic 第一次正式对外建立"能力解锁许可证"体系。')
para('对行业观察者: Model Tiering 是未来两年的核心看点。观察 OpenAI 和 Google 会不会跟进类似分层，监管者会不会把这种分层写进合规模板。')

# ============ Closing ============
h1('结语')
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.line_spacing = 2.0
r = p.add_run('三条带走的判断:\n')
set_font(r, size=12, bold=True, color='B07B1F')

points = [
    '① Opus 4.7 的真正价值是"同等质量减半 token"，不是 benchmark 分数。迁移决策看 xhigh 档位的真实生产流量表现。',
    '② Mythos Preview + Project Glasswing 建立了 AI 发布的新范式(Model Tiering)，未来 OpenAI/Google 大概率会跟进。',
    '③ 差异化训练是 Anthropic 接下来两年的核心护城河——能把一个底座模型切出 N 个合规版本的能力，比"再训一个更强的"更有商业价值。',
]
for t in points:
    p = doc.add_paragraph(t)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.line_spacing = 1.7
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        set_font(run, size=11, color='37352F')

# shading background for closing box
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
for p in doc.paragraphs[-4:]:
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(r'<w:shd {} w:fill="FFF8E6"/>'.format(nsdecls('w')))
    pPr.append(shd)

# references
h2('引用')
refs = [
    '1. Introducing Claude Opus 4.7 — https://www.anthropic.com/news/claude-opus-4-7 (Anthropic, 2026-04-16)',
    '2. Claude Opus 4.7 leads on SWE-bench and agentic reasoning — thenextweb.com',
    '3. Anthropic Releases Claude Mythos Preview — infoq.com',
    '4. Anthropic limits Mythos AI rollout — CNBC 2026-04-07',
    '5. Claude Opus 4.7 in Amazon Bedrock — AWS',
    '6. Anthropic launches Claude Opus 4.7, migration advice — Constellation Research',
]
for r_ in refs:
    p = doc.add_paragraph(r_)
    for run in p.runs:
        set_font(run, size=10, color='787066')
    p.paragraph_format.space_after = Pt(2)

doc.save(OUT)
print('saved:', OUT)
