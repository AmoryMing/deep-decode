#!/usr/bin/env python3
# coding: utf-8
# 用「」替代所有中文双引号，避免 ASCII 嵌套问题

import pathlib
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement as OE

BASE = pathlib.Path(__file__).resolve().parent.parent
PNG = BASE / 'material' / 'pngs'
OUT = BASE / '【AI笔记0421】越追AI风口越被反收割.docx'

doc = Document()
style = doc.styles['Normal']
style.font.name = 'PingFang SC'
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'PingFang SC')
style.font.size = Pt(11)

for s in doc.sections:
    s.top_margin = Cm(2.2); s.bottom_margin = Cm(2.2)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)


def set_cn(r, name='PingFang SC'):
    r.font.name = name
    r._element.rPr.rFonts.set(qn('w:eastAsia'), name)


def heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text); set_cn(r); r.bold = True
    r.font.size = Pt(20 if level == 1 else 16)
    r.font.color.rgb = RGBColor(55, 53, 47)


def para(text, bold=False, italic=False, color=(55, 53, 47), size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.55
    r = p.add_run(text); set_cn(r)
    r.bold = bold; r.italic = italic
    r.font.size = Pt(size); r.font.color.rgb = RGBColor(*color)


def quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(8)
    r = p.add_run('「' + text + '」'); set_cn(r)
    r.italic = True; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(120, 119, 116)


def image(png_name, width_cm=15.5, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(); r.add_picture(str(PNG / png_name), width=Cm(width_cm))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(caption); set_cn(cr)
        cr.font.size = Pt(9); cr.italic = True
        cr.font.color.rgb = RGBColor(120, 119, 116)


def end_box(lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    pPr = p._p.get_or_add_pPr()
    pBdr = OE('w:pBdr')
    for side in ('top', 'bottom'):
        b = OE('w:' + side)
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '12'); b.set(qn('w:color'), 'FFB020')
        pBdr.append(b)
    pPr.append(pBdr)
    shd = OE('w:shd'); shd.set(qn('w:fill'), 'FFF7E0'); pPr.append(shd)
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(line); set_cn(r); r.bold = True
        r.font.size = Pt(11); r.font.color.rgb = RGBColor(55, 53, 47)


# ========== 封面 ==========
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(80)
tr = tp.add_run('越追 AI 风口，越被反收割'); set_cn(tr)
tr.bold = True; tr.font.size = Pt(32); tr.font.color.rgb = RGBColor(55, 53, 47)

sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp.paragraph_format.space_before = Pt(12)
sr = sp.add_run('一份给产品人的「潮汐手册」'); set_cn(sr)
sr.font.size = Pt(16); sr.font.color.rgb = RGBColor(120, 119, 116)

tg = doc.add_paragraph(); tg.alignment = WD_ALIGN_PARAGRAPH.CENTER
tg.paragraph_format.space_before = Pt(40)
tgr = tg.add_run('AI 笔记 · 04/21 · 2026\n深度解读 · credit 见文末'); set_cn(tgr)
tgr.font.size = Pt(12); tgr.font.color.rgb = RGBColor(35, 131, 226)

image('00_cover.png', width_cm=15)
doc.add_page_break()

# ========== 开篇 ==========
para('2026 年第一季度，全球风险投资砸进 AI 的钱是 2420 亿美元。这个数字占当季全球 VC 总额的 80%，相当于 2025 年全年 AI 融资的大半。与此同时，麻省理工最新的企业调研显示——95% 接入 AI 的公司，拿不出可量化的投资回报。')
para('一边是史上最大规模的资本集结，一边是 95% 的回本失败率。两个数字放在一起，就是一个价值几千亿美元的问题——你是那 5%，还是那 95%？', bold=True)
image('01_three_mirrors.png', caption='图 1 · 历史三镜：1840 铁路 / 2000 互联网 / 2026 AI')
para('这个问题有一个反直觉的答案：越追 AI 风口，越容易被时代淘汰。不是劝大家远离 AI，是劝大家换个姿势——别当那只追兔子的农夫，做那棵等兔子撞上来的树。')

# ========== 一 ==========
heading('一、潮汐和浪花——一个看世界的视角')
para('第一个值得拿出来的框架，是把世界上所有变化分成两类：快变量和慢变量。')
para('浪花，就是快变量。海面上一拨一拨打过来，声势浩大，动静十足。今天这家公司发了新模型，明天那家推了新 Agent，后天国家出了新文件。浪花的特点是它们会抢走你的全部注意力。你站在海边盯着看一个小时，会觉得海在剧烈变化。')
para('潮汐，就是慢变量。以几乎看不见的速度，一厘米一厘米改变海岸线的形状。技术革命、人口结构、全球化进程、能源转型——这些才是真正塑造未来的力量。')
para('问题是，浪花会自己跳到你眼前，潮汐不会。所以我们每天花 3 小时刷新闻刷热搜，却几乎没有时间去读一份关于人口结构或某个新技术长期趋势的深度报告。', bold=True)
para('这个视角叫「海拔三千米经济学」（出处见文末）。我们平时活在海平面，周围全是高楼大厦车水马龙。上珠峰太难，但爬到海拔三千米，往上能看到云外的雪峰，往下能看到一整片星罗棋布的城镇。')
quote('贸易摩擦只是树枝，国际关系才是大树；股价涨跌只是浪花，技术革命才是洋流。')
para('这套视角对产品人尤其重要。过去三年，多少产品团队把季度目标订在「追上最新一代模型」上？回头看，那些每个月换一次技术栈的团队，大部分既没跟上技术，也没交付产品。因为他们追的是浪花，不是潮汐。')

# ========== 二 ==========
heading('二、1840 年的一次集体踩坑')
para('第二个判断更扎心：在每一轮技术革命的早期，都伴随着大量的资产价格泡沫。')
para('最经典的案例，是 1840 年代英国的「铁路狂热」。')
para('当时所有人都知道铁路是未来——它能把内陆的煤炭运到港口，把工厂的商品运往全国。贵族、商人、市民、家庭主妇，都把积蓄拿出来买铁路公司的股票。短短几年里，英国注册的铁路公司数量暴涨到上千家。最后结果是，大部分公司破产，无数人血本无归。等这批公司倒干净之后，真正的铁路巨头才诞生——因为这时候基础设施已经铺好，做生意的成本已经足够低。')
para('2000 年的互联网泡沫是同一个剧本。纳斯达克从 2000 年 3 月的 5048 点跌到 2002 年的 1114 点，缩水 78%。无数带着 .com 后缀的公司烧光钱死掉。等这波过去，基础设施沉淀下来，谷歌、亚马逊、eBay 才开始真正崛起。亚马逊当时股价也从 113 美元跌到 5.5 美元——贝佐斯那时候说了一句「股票不等于公司」，业务各项指标都在改善，股价还在跌。')
quote('每一轮技术革命，都是前人栽树，后人乘凉。 —— 何帆')
para('现在再看 2026 年 AI 这一轮。OpenAI 单季融 1220 亿美元，Anthropic 300 亿，xAI 200 亿，Waymo 160 亿——这四家合计 1880 亿美元，占 Q1 2026 全球 VC 总额的 65%。资本往四家头部公司极度集中，其余 95% 接入 AI 的企业拿不出量化回报。Alpine Macro 首席量化策略师 Henry Wu 给了一个冷静判断：我们很可能正处于泡沫形成的早期阶段。Henry Wu 进一步说，泡沫的后期特征是反复剧烈的回调，每一次都让空头高呼顶部已到，紧随其后的 V 型复苏又让早期怀疑者缴械投降，直到市场弥漫着"这次不一样"的论调。')
para('「这次不一样」——这五个字是泡沫破裂的背景音。', italic=True)

# ========== 三 ==========
heading('三、守树策略——一个反英雄主义的策略')
image('02_wait_by_tree.png', caption='图 2 · 守树 = 等 AI 在专业必经之路上撞过来')
para('如果泡沫早晚要破，普通人怎么办？有一个反英雄主义的答案：守株待兔。这不是消极，是有意识的战略耐心。')
quote('不要用你的业务，去挑战别人的专业。要用你的专业，去拥抱新技术的红利。 —— 何帆')
para('你是医生，就好好看病。等 AI 跑过来帮你更高效地看片子、做诊断。你是老师，就好好教书育人，等 AI 过来帮你改作业。你是内容创作者，就好好做内容，等 AI 过来帮你生图、剪视频。')
para('这个策略背后有一个容易被忽视的逻辑：新技术真正改变世界的方式，从来不是把每个领域都推倒重来，而是以接口形式接入现有专业。X 光机没有取代医生，CT、MRI 都没有取代医生——这些技术每一代都「更颠覆」，但最后都是成了医生手里的工具。真正赚到红利的，是每个时代都能把最新工具接进自己工作流的那个医生，不是改行去做 X 光机销售的医生。', bold=True)
para('这就浮现出一个值得命名的现象——「风口反收割」：你本以为要去收割风口，结果被风口反过来收割了时间、金钱、注意力，连自己真正的长板都荒废了。被风口反收割的人，比被 AI 替代的人还多。')

# ========== 四 ==========
heading('四、选择题选手 vs 解答题选手')
image('03_answer_vs_choice.png', caption='图 3 · AI 时代最贵的一件事，是放弃「标准答案思维」')
para('还有另一个扎心的分类，可能是全场最狠的一刀。')
para('我们从小是在「寻找标准答案」的教育里长大的。面前有几百家外卖店，A 便宜两块、B 好评多、C 送饮料，花 30 分钟反复对比终于下单——这就是「最优解思维」。结果是，送到的时候一看，同事那份好像更好吃。')
para('赫伯特·西蒙说过，人是「有限理性」的——只能在有限的信息和时间里找那个「够好了」的选项。那个 80 分的满足解。')
para('2010 年有个小产品叫 Burbn，一股脑塞了几十个功能。上线后没人用，只有「给照片加滤镜」的小功能用的人挺多。创始人一咬牙砍掉 90% 的功能，剩下的那个产品就是 Instagram。它不是最优解，但它是从真实反馈里长出来的满足解。')
para('AI 时代最怕的姿势，是站在外卖点餐的心态面前看技术路线。「Claude 和 GPT 哪个好？」「LangChain 和 CrewAI 哪个强？」——这些问题都默认有标准答案。解答题选手先把要解决的问题搞清楚，再去挑工具。选择题选手先把所有工具研究一遍，再问自己「那我要解决什么呢？」——等他想清楚，时代已经换了两代。', bold=True)

# ========== 五 ==========
heading('五、把手弄脏——从显微镜里看到需求')
para('所以问题就变成了：不追风口，那我追什么？答案是三个字：把手弄脏。')
para('过去二三十年，中国经济吃了两个巨型红利——人口红利让你以低成本造东西，流量红利让你以低成本获客。靠一个模式甚至一个点子就赚到钱，是红利期的特产。现在红利没了。赚钱回归到又苦又累的细节里。')
para('最好用的例子是银发经济。一提到银发经济，大多数人脑子里蹦出来的画面是「平台」。但真正的机会不在平台，在「一个 80 岁的老人晚上起夜时会遇到什么」。')
para('一个叫「适老化改造」的行业正在悄悄起来。做的事很具体：把卫生间门槛拆掉让轮椅能进；把瓷砖换成防滑地砖；按老人身高臂长在马桶和淋浴区装支撑扶手；把刺眼白色顶灯换成光线柔和的暖色灯。每一单都要上门测量、设计、施工，赚的是慢钱、小钱。但这个市场正在爆发。')
para('对产品人：过去的产品思维是「规模先于需求」；今天的产品思维是「显微镜先于放大镜」——先把一个具体用户在某个具体场景里的痛点挖到底，再想怎么复制。', bold=True)

# ========== 六 ==========
heading('六、盲区——这次真的一样吗')
para('一个好的分析必须自带反面论证。让我把最大的几个质疑摆出来。')
para('质疑一：AI 的迭代速度跟铁路、互联网真的能对齐吗？', bold=True)
para('1840 年英国铁路狂热到 1870 年代铁路巨头成形，用了 30 年。2000 年互联网泡沫到 Google 真正成为商业帝国，用了 10 年。这两轮的共同特征是——基础设施物理沉淀需要时间。AI 不一样，它的「基础设施」是算力+数据+模型权重，可以在一个晚上搬家。2023-03 GPT-4 发布到 2026-04 已经有二十多个能力相近的替代模型，压缩到 3 年。如果成熟期从 30 年压缩到 3 年，「守树待兔」里那个等待期就短得多。医生还能守，客服可能已经没树可守了。')
para('质疑二：这次泡沫的结构跟 dot-com 不一样。', bold=True)
para('Alpine Macro 的观察是：目前 AI 相关资本支出大多由企业内部现金流支持，而非大规模举债。OpenAI 的 240 亿美元 ARR、Anthropic 的 300 亿美元 ARR——这是 2000 年任何一家 dot-com 公司都拿不出来的真实收入。结构不同，意味着就算有泡沫，破裂的形态也会不同——更可能是几家中小 AI 公司连环失血、头部公司继续吸钱的分化式回调。')
para('质疑三：没有「专业领域」可守的年轻人怎么办？', bold=True)
para('这个策略暗含一个前提——你已经有十年专业积累。对 22 岁刚毕业的年轻人，守什么？零起点新人反而可能是最适合追早期 AI 赛道的那群人——没有被上一代专业绑住的沉没成本。原论述没回答这一题，这是这个框架的盲区。')

# ========== 七 ==========
heading('七、电影院策略——看不清就买张票站门口')
para('针对看不清方向的情况，还有一个很小但很实用的策略，叫「电影院策略」。如果你对某个机会既不能完全错过，又不敢下重注——那就先买张票进场，但站在门口的位置看一会儿。看到剧情好再入座；看到不对劲立刻走人。')
para('具体到 AI 场景——你不必辞职去创业公司，但可以用周末时间跑两个 POC；你不必为公司买一整套企业版，但可以让一个小团队用闲置预算跑一个试点；你不必学一整套 LangChain 架构，但可以用 Claude Code 搭几个小工具解决自己工作流里的三五个真问题。', bold=True)
para('这个策略的本质是保持参与度，同时严格控制沉没成本。')

# ========== 八 ==========
heading('八、对产品人意味着什么')
image('04_reverse_harvest.png', caption='图 4 · 被风口反收割——焦虑税的去向')
para('把这整套框架的信号折回到一个产品人的 2026 年：')
para('第一，停止每周追新工具的惯性。每天 15 分钟读一份人口结构或技术长期趋势的深度报告，比每天 3 小时刷 AI 推特有用一百倍。', bold=True)
para('第二，回到自己领域深挖。如果你是做零售 PM，把 AI 当做接入工具，想清楚它能怎么切进你的选品、定价、会员三个现有环节。', bold=True)
para('第三，从选择题选手变成解答题选手。遇到技术选型问题，先把要解决的业务问题描述到每个人都能复述一遍，再选工具。', bold=True)
para('第四，把手弄脏。不要相信任何只在 PPT 上推演的 AI 产品策略。做一个 100 人以内的小闭环跑通，再考虑扩。', bold=True)
para('第五，电影院策略。不要 all in，也不要完全旁观。小额、多点、快速进出。', bold=True)

para('')
quote('历史的进步从来不是线性的，而是在试错与妥协中蹒跚前行。 —— 何帆《大局观》')
para('线性是泡沫期的错觉。蹒跚才是长期的真相。')

# ========== 结语 ==========
end_box([
    '★ 三条启示',
    '1. 分清潮汐和浪花 —— 每天 15 分钟深度报告，比每天 3 小时热点刷屏，收益高一百倍。',
    '2. 别挑战别人专业 —— 守在自己十年积累的专业里，用专业接 AI，比用业务挑战 AI 存活率高。',
    '3. 做解答题选手 —— 先定义业务问题，再挑工具。顺序反了，9 成返工。',
])

heading('引用')
refs = [
    '① 何帆《越追 AI 风口，越容易被时代淘汰》B 站视频 https://b23.tv/Ana69zK',
    '② 何帆《大局观：真实世界中的经济学思维（新版）》2025',
    '③ 《2026，看清变量，保持耐心》虎嗅网 2026-01',
    '④ Crunchbase《Q1 2026 Shatters Venture Funding Records》',
    '⑤ Remio《AI Venture Capital Surges to 242B in Q1 2026》',
    '⑥ 证券时报《2026 全球市场前瞻：美股 AI 泡沫初期》',
    '⑦ 上海金融与发展实验室《AI 泡沫？英伟达估值分析报告》2026-01',
    '⑧ 何帆 财新博客《在慢变量中寻找小趋势》',
]
for r in refs:
    para(r, size=9, color=(120, 119, 116))

doc.save(str(OUT))
print('OK:', OUT.name)
