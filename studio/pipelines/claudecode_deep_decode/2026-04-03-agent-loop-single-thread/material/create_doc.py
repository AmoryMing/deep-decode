#!/usr/bin/env python3
"""生成【AI笔记0407】单线程Agent Loop.docx"""
import os, sys
sys.path.insert(0, '/tmp')

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
except ImportError:
    os.system('pip3 install python-docx -q')
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.abspath(__file__))
PNGS = os.path.join(BASE, 'pngs')
OUT_DIR = os.path.dirname(BASE)

doc = Document()

# -- 页面设置 --
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

style = doc.styles['Normal']
style.font.name = 'Songti SC'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Songti SC')

def add_title(text, size=22, bold=True, color=RGBColor(0x1A, 0x1A, 0x1A)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    run.font.name = 'Heiti SC'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Heiti SC')
    return p

def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = 'Heiti SC'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Heiti SC')
    return p

def add_para(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = Pt(22)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bold_para(prefix, bold_text, suffix=''):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(22)
    if prefix:
        p.add_run(prefix)
    run = p.add_run(bold_text)
    run.bold = True
    run.font.color.rgb = RGBColor(0xC4, 0x7B, 0x2B)
    if suffix:
        p.add_run(suffix)
    return p

def add_quote(en_text, cn_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(en_text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x5D, 0x4E, 0x37)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1)
    run2 = p2.add_run(f'-- {cn_text}')
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x5D, 0x4E, 0x37)

def add_img(name, width=5.5):
    png_path = os.path.join(PNGS, f'{name}.png')
    if os.path.exists(png_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(png_path, width=Inches(width))
    else:
        print(f'[WARN] PNG not found: {png_path}')

# ============================================================
# 封面
# ============================================================
add_title('Claude Code 源码拆解系列 #5', size=12, bold=False, color=RGBColor(0x78, 0x77, 0x74))
doc.add_paragraph()
add_title('单线程 Agent Loop', size=26)
add_title('为什么 Claude Code 不用 Swarm', size=16, bold=False, color=RGBColor(0x78, 0x77, 0x74))
doc.add_paragraph()
add_img('00_系列封面', width=5.0)
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('多来源综合 | 2026-04-07 | 整理：AI Force')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x78, 0x77, 0x74)

# 分页
doc.add_page_break()

# ============================================================
# 正文
# ============================================================
add_bold_para('', 'while (true)', ' —— 51 万行代码的 AI 编程工具，心脏只有这四个单词。')
add_para('2024 年 10 月，OpenAI 开源了 Swarm，一个多 Agent 编排框架，核心卖点是 Agent 之间的 handoff 和 routine。社区一片叫好：多 Agent 是未来。2026 年 3 月，Claude Code 源码泄露，51 万行 TypeScript 摊在阳光下。所有人都去找它的多 Agent 系统——确实找到了三套。但翻遍 1900 个文件后，一个更有趣的事实浮出水面：Claude Code 的默认模式是单线程。一个 while(true) 循环，一条消息历史，模型自己决定下一步做什么。')
add_para('三套多 Agent 系统全部是 feature-flagged 的可选层。Anthropic 把多 Agent 当作特定场景的解决方案，而非通用范式。')

add_heading('一条 while 循环的一生', level=1)
add_para('打开 query.ts，核心心跳就在第 307 行。注释 no-constant-condition 说明连 linter 都觉得这写法不对劲。但这恰好揭示了 Claude Code 的架构哲学：不是模型服从编排器的调度，而是模型自己就是调度器。')
add_para('这个循环每一轮做四件事：准备上下文 → 调 Claude API → 执行工具 → 判断是否继续。源码中有两个安全阀：maxTurns（最大循环轮次）和 maxBudgetUsd（预算上限）。Agent 跑飞了怎么办？设个上限就好。')
add_bold_para('整个设计的核心假设：', '模型足够聪明，可以充当自己的调度器。', '不需要外部编排框架替它决定"下一步做什么"，不需要 DAG 来预定义执行流程。')
add_bold_para('姑且给这个设计哲学一个名字：', '"模型即调度器"（Model-as-Scheduler）。')

add_img('01_agent_loop_vs_swarm', width=5.5)

add_heading('不是没有并发，是在正确的层级做并发', level=1)
add_para('单线程三个字容易让人误解为不并发。恰好相反，Claude Code 在单线程内实现了精细的并发控制。')
add_para('源码中有个关键类 StreamingToolExecutor，它把工具分成两类——并发安全的（读文件、搜索、glob）可以同时跑；需要串行的（写文件、执行 bash）必须排队。并发安全不是固定标签，而是根据实际参数动态判断。同一个 BashTool，ls 是并发安全的，rm -rf 不是。')
add_para('更激进的是投机执行。在 Auto 权限模式下，模型还在流式输出工具调用参数时，StreamingToolExecutor 已经开始做安全分类检查了。等模型输出完成时，分类器结果可能已经就绪，用户零等待。')
add_bold_para('准确的描述是：', '主循环单线程保证全局一致性，工具层并发最大化局部吞吐量。')

add_img('02_concurrency_ladder', width=5.5)

add_heading('单线程的经济学：prompt cache 红利', level=1)
add_para('选择单线程不只是架构偏好，背后有一笔精确的经济账。Sebastian Raschka 提出核心论点：harness 比 model 更重要。其中最关键的 harness 技术就是 prompt cache——Anthropic API 对请求前缀相同的调用，缓存命中部分的计费只有正常价格的 10%。')
add_para('单线程架构天然适合 prompt cache。每次 while 循环迭代时，系统提示词、CLAUDE.md、memory、工具定义、消息历史的前缀部分完全相同。LMCache 团队的分析显示，Claude Code 的 prompt cache 命中率高达 92%。')
add_bold_para('', '"单线程红利"', '：单线程不只是架构简洁，它是 prompt cache 经济模型的最优解。一旦切换到多 Agent 默认范式，这笔红利就消失了。')

add_heading('三套多 Agent 系统：为什么都不是默认', level=1)
add_heading('Fork 子 Agent：缓存优先的轻量并行', level=2)
add_para('最轻量的方案。子 Agent 共享父 Agent 的 prompt cache 前缀，独立执行后返回。不是独立 Agent，更像主循环的"分身术"——做完就收回，不维护独立状态。')

add_heading('Swarm/Teammate：文件邮箱团队系统', level=2)
add_para('基于文件邮箱的 Agent 团队系统。每个 teammate 有一个 JSON 邮箱文件，Agent 之间通过读写邮箱通信。可运行在同进程模式（AsyncLocalStorage 隔离）或分屏模式（tmux/iTerm2 独立进程）。')

add_heading('Coordinator：不许偷懒的项目经理', level=2)
add_para('最重的多 Agent 模式。Claude 变成只有 3 个工具的项目经理：Agent（派活）、SendMessage（追加指令）、TaskStop（叫停）。定义了四阶段工作流：Research → Synthesis → Implementation → Verification。')

add_quote('"Never write \'based on your findings\' or \'based on the research.\' These phrases delegate understanding to the worker instead of doing it yourself."', '不要写"根据你的发现"——这是把理解委派给了 worker。—— Coordinator System Prompt')

add_bold_para('姑且叫做', '"认知不可委派原则"（Non-Delegable Understanding）', '：在 Agent 编排中，理解问题这一步永远不能外包给子 Agent。可以委派执行，但不能委派理解。')

add_img('03_coordinator_workflow', width=5.5)

add_heading('对 AI 产品实践者意味着什么', level=1)
add_bold_para('1. ', '默认单线程，按需多 Agent。', ' 不要因为"多 Agent 听起来更先进"就默认选择多 Agent 架构。先问自己：任务本质上是串行的还是可并行的？')
add_bold_para('2. ', '如果做多 Agent，先想清楚"理解"在哪个层级发生。', ' Coordinator 铁律指出了多 Agent 系统中最容易犯的错误：把理解问题这一步也委派出去。')
add_bold_para('3. ', '关注 prompt cache 经济学。', ' 单线程架构天然适合 prompt cache（92% 命中率）。如果切换到多 Agent，要算清楚 cache 命中率下降带来的成本增加。')

# ============================================================
# 结语
# ============================================================
doc.add_page_break()
add_heading('本期关键词', level=1)
add_bold_para('', 'Model-as-Scheduler（模型即调度器）', ' —— 不用外部编排框架，让模型自己决定下一步做什么。Claude Code 的 while(true) 循环是这个理念的极致体现。')
add_bold_para('', 'StreamingToolExecutor（流式工具执行器）', ' —— 在单线程内实现工具级并发。读操作并行、写操作串行，并发安全性根据实际输入动态判断。')
add_bold_para('', 'Non-Delegable Understanding（认知不可委派）', ' —— 理解问题这一步永远不能外包给子 Agent。可以委派执行，但不能委派理解。管理哲学的代码化。')
add_bold_para('', '单线程红利', ' —— prompt cache 命中率 92%。一旦切换到多 Agent 默认模式，前缀重叠率下降，这笔红利消失。')

# -- 保存 --
out_path = os.path.join(OUT_DIR, '【AI笔记0407】单线程Agent Loop.docx')
doc.save(out_path)
print(f'Done: {out_path} ({os.path.getsize(out_path)} bytes)')
