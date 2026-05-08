#!/usr/bin/env python3
# 生成【AI笔记0407】Undercover Mode Anthropic员工的隐身衣.docx
import os, sys
sys.path.insert(0, '/tmp')
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    os.system('pip3 install python-docx -q')
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.abspath(__file__))
PNGS = os.path.join(BASE, 'pngs')
OUT_DIR = os.path.dirname(BASE)
doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2); s.left_margin = Cm(2.8); s.right_margin = Cm(2.8)
style = doc.styles['Normal']; style.font.name = 'Songti SC'; style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33,0x33,0x33); style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Songti SC')

def title(t, sz=22):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.font.size = Pt(sz); r.bold = True; r.font.color.rgb = RGBColor(0x1A,0x1A,0x1A)
    r.font.name = 'Heiti SC'; r.element.rPr.rFonts.set(qn('w:eastAsia'), 'Heiti SC')
def sub(t):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x66,0x66,0x66)
def h(t, lv=1):
    p = doc.add_heading(t, level=lv)
    for r in p.runs: r.font.name = 'Heiti SC'; r.element.rPr.rFonts.set(qn('w:eastAsia'), 'Heiti SC')
def para(t):
    p = doc.add_paragraph(t); p.paragraph_format.line_spacing = Pt(22); p.paragraph_format.space_after = Pt(6)
def img(name, w=5.8):
    path = os.path.join(PNGS, name)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(w))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
def kw(term, desc):
    p = doc.add_paragraph(); p.paragraph_format.line_spacing = Pt(20)
    r = p.add_run(term); r.bold = True; r.font.color.rgb = RGBColor(0xC4,0x7B,0x2B)
    p.add_run(f' -- {desc}')

# 封面
doc.add_paragraph()
title('Undercover Mode', sz=28)
title('Anthropic 员工的隐身衣', sz=18)
doc.add_paragraph()
sub('Claude Code 源码深度拆解 #6')
sub('2026-04-07')
img('00_系列封面.png', w=4.5)
doc.add_page_break()

# 正文
para('"93% 3-shotted by claude-opus-4-5, 2 memories recalled." 这行字出现在 Anthropic 内部仓库的 PR 描述里。但你永远不会在任何公开仓库里看到它——因为一个叫 Undercover Mode 的系统自动抹掉了它。')
para('源码注释里有一句话比任何功能描述都更有信息量："There is NO force-OFF." 没有强制关闭选项。无法确认是内部仓库时，默认隐身。')

h('一套源码，两张面孔')
para('当 Anthropic 员工在公共仓库工作时，系统自动注入 CRITICAL 级别提示词，禁止一切 AI 归属信息。措辞像间谍小说："Do not blow your cover."')
img('01_五层纵深防御.png')
para('五层纵深防御：构建时死代码消除 → 运行时 USER_TYPE 检查 → 仓库白名单比对 → CRITICAL 提示词注入 → 归属信息返回空字符串。每一层都是上一层的兜底。')

h('21 个仓库背后的 Anthropic')
img('03_内部仓库白名单.png')
para('内部仓库白名单列了 21 个私有仓库，揭示了 Anthropic 的组织结构：核心产品（CLI/Web/移动端）、基础设施（Terraform/K8s）、工程工具链、数据平台、甚至 AI 辅助招聘。关键：anthropics/claude-code 是公共仓库，不在白名单上。')

h('字符级贡献度追踪')
para('同一套源码里还有一套完整的代码归属追踪系统：每次 Claude 编辑文件，记录修改字符数。到 PR 时汇总百分比 + N-shot 计数 + 记忆调用次数。这套系统只在内部仓库启用。')

h('归属悖论')
img('02_归属悖论.png')
para('同一套源码同时在做两件矛盾的事：一边抹掉 AI 身份，一边精确追踪 AI 贡献。不矛盾——隐身解决今天的商业机密问题，追踪为明天的合规法规做准备。EU AI Act 2026 年 8 月 2 日起要求标注 AI 生成内容。')

h('对 AI 从业者意味着什么')
para('AI 归属不是技术问题，是治理问题，窗口期正在关闭。做 AI 工具的现在就建归属追踪系统；开源维护者需要明确 AI 贡献政策；企业合规负责人需要知道 AI 贡献了多少代码。')

doc.add_page_break()
h('本期关键词')
kw('Undercover Mode', '当 Anthropic 员工在公共仓库工作时自动激活的隐身系统，抹掉所有 AI 归属信息')
kw('DCO（Developer Certificate of Origin）', '开源贡献者证书，AI 隐身提交与 DCO 存在冲突')
kw('Attribution Tracking', '字符级代码贡献追踪，记录 AI vs 人类的贡献比例')
kw('Feature Flag 死代码消除', '编译时把未启用功能物理删除，Undercover 在外部构建中不存在')
kw('EU AI Act Article 50', '2026年8月2日起要求标注 AI 生成内容')

h('引用')
for ref in ['1. Claude Code v2.1.88 泄露源码 -- 一手素材',
    '2. Ars Technica -- Claude Code 泄露报道',
    '3. WaveSpeed AI -- Undercover Mode 深度分析 + DCO 讨论',
    '4. HN 讨论 -- 开发者社区观点碰撞',
    '5. EU AI Act Article 50 -- 2026.08.02 标注义务',
    '6. AI Code Share 行业数据 -- 2026 AI 代码占比',
    '7. EU Commission Code of Practice -- 标注实践准则']: para(ref)

out = os.path.join(OUT_DIR, '【AI笔记0407】Undercover Mode Anthropic员工的隐身衣.docx')
doc.save(out); print(f'OK: {out}')
