#!/usr/bin/env python3
"""
Claude Code 源码泄露全景 — Word 文档生成（python-docx 版）
运行：python3 material/create_doc.py
"""
import os, sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SCRIPT_DIR = Path(__file__).parent
PNGS_DIR = SCRIPT_DIR / "pngs"
OUT_DIR = SCRIPT_DIR.parent

# 颜色
C_TITLE = RGBColor(0x1A, 0x1A, 0x1A)
C_ACCENT = RGBColor(0xC4, 0x7B, 0x2B)
C_BODY = RGBColor(0x33, 0x33, 0x33)
C_LIGHT = RGBColor(0x66, 0x66, 0x66)
C_HI = RGBColor(0xD2, 0x69, 0x1E)

doc = Document()

# 默认样式
style = doc.styles['Normal']
style.font.name = 'Songti SC'
style.font.size = Pt(10.5)
style.font.color.rgb = C_BODY
style.paragraph_format.line_spacing = 1.5

# ============ 辅助函数 ============

def add_img(name, width_inches=5.5):
    png = PNGS_DIR / f"{name}.png"
    if not png.exists():
        print(f"[警告] PNG 不存在：{png}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.add_run().add_picture(str(png), width=Inches(width_inches))

def add_h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = C_TITLE
    run.font.name = 'Heiti SC'

def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = C_ACCENT
    run.font.name = 'Heiti SC'

def add_p(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = C_BODY

def add_pb(before, bold_text, after=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(before)
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = C_BODY
    r2 = p.add_run(bold_text)
    r2.font.size = Pt(10.5)
    r2.font.bold = True
    r2.font.color.rgb = C_HI
    if after:
        r3 = p.add_run(after)
        r3.font.size = Pt(10.5)
        r3.font.color.rgb = C_BODY

def add_li(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.8)
    r0 = p.add_run("  \u2022  ")
    r0.font.color.rgb = C_ACCENT
    r1 = p.add_run(text)
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = C_BODY

def add_quote(en, cn):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(1)
    r1 = p.add_run(en)
    r1.font.italic = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0x5D, 0x4E, 0x37)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1)
    p2.paragraph_format.space_after = Pt(8)
    r2 = p2.add_run(f"-- {cn}")
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = C_LIGHT

def add_divider():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("\u2022  \u2022  \u2022")
    run.font.color.rgb = RGBColor(0xE8, 0xD5, 0xB5)
    run.font.size = Pt(10)

# ============ 封面 ============

add_img("00_系列封面", 5.8)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Claude Code 源码拆解系列 01/34  \u00b7  AI Force 前沿探索")
r.font.size = Pt(8.5)
r.font.color.rgb = C_LIGHT

doc.add_page_break()

# ============ 正文 ============

add_h2("开篇")

add_p("59.8 MB。这是一个 source map 文件的大小。")

add_p("2026 年 3 月 31 日，Solayer Labs 研究员 Chaofan Shou 在 npm 包 @anthropic-ai/claude-code 的 v2.1.88 中发现了这个异常文件。解压后是 1,902 个 TypeScript 源码文件，512,000 行代码\u2014\u2014Anthropic 最赚钱的产品之一，Claude Code 的完整实现，就这么躺在了全世界每一个开发者的 node_modules 里。")

add_p("GitHub 镜像仓库在 48 小时内拿到了 30,000 颗星、40,200 次 fork（CyberNews 数据）。Anthropic 官方回应只有两个词：\u300chunan error\u300d（The Guardian 报道）。")

add_p("但这不是一篇关于\u201c某公司出了安全事故\u201d的文章。事故本身只是一扇窗。透过这 1,902 个文件，能看到的是：当今最复杂的 AI 编程助手到底长什么样，以及\u2014\u2014它的竞争力到底在哪里。")

add_divider()

# ====== 01 ======
add_h1("01  同一个坑，摔了两次")

add_p("先说让人不舒服的部分。")

add_p("这不是 Anthropic 第一次因为 source map 泄露源码。2025 年 2 月，v0.2.8 版本就出过一次\u2014\u2014当时的代码量小得多，base64 反混淆后能看到的信息有限，行业没当回事。")

add_pb("16 个月后，同一类错误以更大规模重演。dev.to 作者 Varshith Hegde 在复盘文章中提到了一个关键细节：", "Anthropic 收购了 Bun（Claude Code 的 JavaScript runtime），而 Bun 的 source map 生成 bug 在事发 20 天前就被社区报告了，但没有修复。")

add_img("06_泄露时间线")

add_pb("这暴露了什么？不是某个工程师的疏忽\u2014\u2014这种说法太轻了。数据表明这是 ", "CI/CD 流水线的系统性缺陷", "：从代码构建到 npm 发布的链路上，没有一个环节检测到一个 59.8MB 的异常文件混进了分发包。")

add_p("Varshith Hegde 还提出了一个更尖锐的问题：事发日是 3 月 31 日，愚人节前夜；源码中发现的 Buddy 宠物系统原定 4 月 1\u20147 日上线。真的是意外？还是一次精心设计的 PR 事件？")

add_pb("Anthropic 不太可能故意泄露自己的安全架构细节和未发布功能。但这个质疑之所以值得记录，是因为它指向了一个更深层的问题：", "当一家以安全著称的公司连续犯同类错误，行业对它的信任评估需要重新校准。")

add_divider()

# ====== 02 ======
add_h1("02  六层自研：不是\u300c模型加壳\u300d")

add_p("拆开这 512,000 行代码，第一个感受是：这不是大多数人想象中的\u201c聊天机器人加个终端壳\u201d。")

add_p("Engineer\u2019s Codex 的分析最直观\u2014\u2014整个代码库按功能分为 7 大类、40 个工具模块、184 个文件，仅工具系统就有约 50,800 行代码（Randal Olson 统计）。而这只是冰山水面上的部分。")

add_img("01_六层架构")

add_p("从底层到顶层，整套系统可以归纳为六层架构：")

add_pb("L1 运行时层。", "Bun 作为 JavaScript runtime，", "负责启动和底层 I/O。这是 Anthropic 收购 Bun 后的直接产物\u2014\u2014自己控制运行时，不受 Node.js 生态的节奏约束。")

add_pb("L2 核心引擎层。", "QueryEngine 是整个系统的大脑中枢，", "仅这一个模块就有 13,000 行代码，负责对话编排、Prompt 组装、工具调度、上下文压缩（Snip 机制）和 Token 追踪。PromptLayer 的技术分析用\u201cmaster agent loop\u201d来描述这个引擎\u2014\u2014单线程循环，不用 Swarm 多 Agent 架构，所有决策在一个循环里完成。")

add_p("这个设计选择值得停下来说一下。行业里有一股明显的趋势是用多 Agent 编排（如 AutoGen、CrewAI）来处理复杂任务。但 Claude Code 选择了相反的路：单线程、单循环、所有工具平等接入。这意味着它赌的是单个模型足够强，不需要多个弱模型协调。")

add_pb("L3 工具层。", "45 个内置工具，每个工具都是一等公民", "\u2014\u2014有独立的描述、参数 schema、权限声明。值得注意的是 MCP（Model Context Protocol）的接入\u2014\u2014这意味着 Claude Code 的工具系统不是封闭的，任何人可以通过 MCP 协议扩展它的能力边界。")

add_pb("L5 记忆层。", "四分类记忆系统（user/feedback/project/reference），", "加上 Snip 压缩机制处理长对话的上下文衰减。LMCache 的分析提到了一个惊人的数据：92% 的缓存复用率\u2014\u2014意味着 Claude Code 在多轮对话中，有超过九成的 prompt 前缀是从缓存读取而非重新计算的。")

add_img("04_记忆系统")

add_pb("L6 终端 UI 层。", "用 Ink（React for CLI）搭建，80+ 个组件文件。", "187 个 spinner 动词、Vim 模态编辑器、全键盘操作\u2014\u2014目标用户画像非常明确：重度终端用户。")

add_pb("这六层架构有一个共同特征：", "核心链路零外部依赖。", "不用 LangChain，不用 AutoGen，不用 Semantic Kernel。Anthropic 选择全栈自研的成本极高，但换来的是对每一层行为的完全控制。")

add_pb("这里可以给一个更结构化的判断：把这个现象命名为\u300c自研溢价\u300d。", "当产品的核心竞争力是\u201c可靠性\u201d时，每一层外部依赖都是一个不可控的风险源。", "自研的成本在前期是负担，但在产品成熟期会变成护城河。")

add_divider()

# ====== 03 ======
add_h1("03  双 AI 制衡：安全不是功能，是地基")

add_p("如果只看一层，看安全层。")

add_pb("Claude Code 的安全架构不是\u201c在执行前加一道审核\u201d这么简单。它本质上是一个 ", "双 AI 制衡系统", "：主 AI（Opus/Sonnet）负责执行任务，一个独立的 Sonnet 分类器（Temperature=0，确保确定性输出）负责判断每一步操作的风险等级。")

add_img("02_双AI安全架构")

add_p("每一条 bash 命令在执行前，要过四道门：")
add_li("bashClassifier\u2014\u2014用 LLM 判断命令的风险类别（只读/写入/破坏性/网络访问）")
add_li("yoloClassifier\u2014\u2014在用户开启\u201c自动模式\u201d时做二次确认")
add_li("权限网关\u2014\u2014根据用户设定的权限等级（从 default 到 auto，共 6 级）决定放行、确认或拒绝")
add_li("熔断机制\u2014\u2014累计异常达到阈值后自动中断会话")

add_quote("Don\u2019t give agent access to everything just because you\u2019re lazy.", "CrowdStrike CTO，VentureBeat 引用")

add_pb("这句话精确描述了 Claude Code 安全架构的设计哲学\u2014\u2014", "信任不是开关，是阶梯。", "6 级权限不是 6 个选项，而是信任从零到最大的连续光谱。")

add_pb("GitGuardian 的数据值得警惕：", "Claude Code 辅助的代码提交中，密钥泄露率为 3.2%，是行业平均水平 1.5% 的两倍多。")

add_pb("这里浮现出一个值得命名的现象，称之为\u300c信任传递悖论\u300d：", "你越信任工具、越频繁地让它自动执行，它帮你犯错的概率就越高。")

add_divider()

# ====== 04 ======
add_h1("04  四个隐藏功能：泄露的不是代码，是路线图")

add_p("源码泄露最让竞品工程师兴奋的部分，不是已上线的功能，而是 Feature Flags 后面藏着的未发布功能。")

add_img("03_隐藏功能矩阵")

add_pb("KAIROS\u2014\u20147x24 后台 Daemon。", "源码注释显示原计划 4 月 1 日上线。", "Daemon 意味着 Claude Code 可以在用户不操作时持续监控代码仓库、自动处理 CI/CD 失败、主动发起 PR review。如果上线，Claude Code 就不再是\u201c你问它答\u201d的工具，而是一个始终在线的工程搭档。")

add_pb("Buddy 系统\u2014\u201418 个物种、5 级稀有度、1% 闪光率、RPG 属性。", "这是一个完整的电子宠物系统，藏在 AI 编码工具里。", "本质上是用户留存机制\u2014\u2014创造一个与编码任务无关的\u201c每天打开\u201d理由。")

add_pb("Undercover Mode\u2014\u2014约 90 行代码，功能是", "抹除所有 AI 辅助痕迹。", "讽刺的是，这个用来\u201c反泄露\u201d的功能本身被泄露了。")

add_pb("Anti-Distillation\u2014\u2014最值得关注。fake_tools 机制会在检测到输出可能被用于训练竞品模型时，", "注入虚假的 tool definitions，本质上是数据投毒。")

add_pb("四个功能放在一起看，能拼出一张清晰的产品路线图：", "KAIROS 是使用深度的延伸，Buddy 是使用频度的延伸，Undercover 是企业客户需求，Anti-Distillation 是竞争防御。", "每一个都是商业目标的直接映射。")

add_divider()

# ====== 05 ======
add_h1("05  屎山也在：3167 行的单函数")

add_p("在赞叹架构精密的同时，源码也暴露了工程质量的另一面。")

add_pb("Engineer\u2019s Codex 提到了一个让工程师倒吸凉气的数据：print.ts 文件总计 5,594 行，其中一个单独的函数就有 ", "3,167 行", "，嵌套深度达到 12 层。")

add_pb("但更有趣的是另一个发现：代码中充满了", "写给 AI 而不是人类的注释。", "\u300cLLM-oriented comments throughout, written for AI agents not humans.\u300d这意味着 Claude Code 的代码库本身就是被 AI 辅助编写和维护的\u2014\u2014它在用自己来开发自己。")

add_pb("Gartner 的判断：Claude Code 的源码被评估为 90% 由 AI 生成，这在美国版权法下意味着", "大部分代码可能不受 IP 保护。")

add_img("05_遥测三通道")

add_p("源码还揭示了完整的遥测架构：Datadog（实时监控）+ BigQuery（数据分析）+ GrowthBook（A/B 测试），三个通道覆盖了从系统健康到用户行为到功能实验的全链路。")

add_divider()

# ====== 06 ======
add_h1("06  护城河评估：被削弱的和没被动摇的")

add_p("最后回到核心问题：源码泄露后，Claude Code 的竞争优势还在吗？")

add_img("07_护城河矩阵")

add_h2("被削弱的")
add_li("安全架构的隐蔽性。四道安全门的具体实现、分类器判断逻辑、熔断阈值，以前是黑箱，现在是白箱。")
add_li("未发布功能的先发优势。KAIROS、Buddy 的设计思路被公开，竞品可以提前布局。")
add_li("代码实现的独占性。竞品工程师可以直接参考 QueryEngine、Snip、权限系统的实现。")

add_h2("没被动摇的")
add_li("工程判断力。知道\u201c怎么做\u201d不等于知道\u201c为什么这么做\u201d。设计取舍存在于团队认知中，不在源码里。")
add_li("模型-工具协同调优。工具描述、Prompt 模板、参数默认值，都针对 Claude 模型特性调优。换模型底座，效果不同。")
add_li("迭代速度。源码是某个时间点的快照。以 Anthropic 当前的迭代速度，公开的代码很快就会过时。")
add_li("92% 缓存复用率。这个数据背后是 prompt engineering 和缓存策略的长期打磨，不是看了代码就能复制的。")

add_pb("给出一个更凝练的判断\u2014\u2014称之为\u300c护城河迁移\u300d：", "源码泄露把 Claude Code 的护城河从\u201c代码壁垒\u201d迁移到了\u201c能力壁垒\u201d。", "代码壁垒是静态的、可复制的；能力壁垒是动态的、需要持续投入的。从长期看，这反而可能是更健康的竞争态势。")

add_divider()

# ====== 结语 ======
add_h2("对 AI 从业者和实践者意味着什么")

add_p("三条可执行的判断：")

add_pb("1. ", "Harness > Model 已被实证。", "围绕模型的工程体系决定产品天花板。与其对比哪个模型 benchmark 高 2%，不如打磨 Agent loop、安全机制和上下文管理。")

add_pb("2. ", "供应链安全需要从\u201c做了\u201d升级到\u201c持续做\u201d。", "产品安全和发布安全是两个独立的问题。Source map、env 文件的排除检查应该是 CI/CD 的必选项。")

add_pb("3. ", "AI 生成代码的 IP 归属是悬而未决的风险。", "关键的架构决策和核心算法应该保留人工编写的记录，至少在法律框架明确之前。")

add_divider()

p = doc.add_paragraph()
r = p.add_run("Claude Code 源码拆解系列 01/34 \u00b7 AI Force 前沿探索")
r.font.size = Pt(9)
r.font.color.rgb = C_LIGHT

p2 = doc.add_paragraph()
r2 = p2.add_run("信源：CyberNews / VentureBeat / Engineer\u2019s Codex / Zscaler / The Guardian / dev.to / Randal Olson / APIYI / PromptLayer / LMCache")
r2.font.size = Pt(8)
r2.font.color.rgb = C_LIGHT

# ============ 保存 ============
out_path = OUT_DIR / "\u3010AI\u7b14\u8bb00406\u30111902\u4e2a\u6587\u4ef6\u91cc\u85cf\u4e86\u4ec0\u4e48.docx"
doc.save(str(out_path))
print(f"docx: {out_path}")
