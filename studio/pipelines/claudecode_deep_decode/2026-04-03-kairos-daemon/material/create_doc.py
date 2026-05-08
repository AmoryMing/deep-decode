#!/usr/bin/env python3
# 生成【AI笔记0407】KAIROS 7x24后台Daemon的野心.docx
# 功能：封面 + 图文穿插正文 + 关键词表 + 引用列表
import os, sys
sys.path.insert(0, '/tmp')

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    os.system('pip3 install python-docx -q')
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.abspath(__file__))
PNGS = os.path.join(BASE, 'pngs')
OUT_DIR = os.path.dirname(BASE)

doc = Document()

# -- 页面设置 --
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

style = doc.styles['Normal']
style.font.name = 'Songti SC'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Songti SC')

# -- 工具函数 --
def add_title(text, size=22):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    run.font.name = 'Heiti SC'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Heiti SC')

def add_subtitle(text, size=13):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = 'Heiti SC'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Heiti SC')

def add_para(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = Pt(22)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x5D, 0x4E, 0x37)

def add_image(name, width=5.8):
    img_path = os.path.join(PNGS, name)
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_keyword(term, desc):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(20)
    run = p.add_run(term)
    run.bold = True
    run.font.color.rgb = RGBColor(0xC4, 0x7B, 0x2B)
    p.add_run(f' -- {desc}')

# ============================================================
#  封面
# ============================================================
doc.add_paragraph()  # spacing
add_title('KAIROS', size=28)
add_title('7x24 后台 Daemon 的野心', size=18)
doc.add_paragraph()
add_subtitle('Claude Code 源码深度拆解 #5')
add_subtitle('2026-04-07')
add_image('00_系列封面.png', width=4.5)
doc.add_page_break()

# ============================================================
#  正文
# ============================================================
add_quote('"Each wake-up costs an API call, but the prompt cache expires after 5 minutes of inactivity -- balance accordingly."')
add_para('')
add_para('这句话藏在 Claude Code 源码的 SleepTool 提示词里。一个 AI 编程助手的内部提示词，在教 AI 怎么省钱。不是教它写更好的代码，不是教它理解用户意图——是教它算账：醒来一次花多少钱，睡太久缓存过期又要多花多少钱。')
add_para('这条提示词指向一个事实：Anthropic 在源码里藏了一个完整的后台守护进程系统，代号 KAIROS。它被 feature flag 锁死，从未出现在任何公开版本里，但在源码中被引用超过 150 次，关联 75 个文件。这不是还没做完的功能，是一套完整的、已通过编译的基础设施，只差一个开关。')
add_para('KAIROS 来自古希腊语，意思是"恰当的时机"。在 Anthropic 的语境下含义更直白：让 AI 拥有自己的时间感——知道什么时候该醒来干活，什么时候该闭嘴睡觉，什么时候该整理记忆。')

add_heading('从传话筒到守夜人')
add_para('今天所有主流 AI 编程工具——Claude Code、Cursor、GitHub Copilot、Codex——都运行在同一个模式下：你问，它答。你不问，它就停。这个模式叫 request-response。')
add_para('request-response 的问题不是回答得不好，而是它把人变成了瓶颈。一个开发者每天实际和 AI 交互大概 2-3 小时。剩下的 21 小时，AI 什么都没干。KAIROS 要解决的就是这 21 小时。')

add_heading('五个子系统，一个守护进程')
add_image('02_五大子系统全景.png')
add_para('KAIROS 由五个子系统组成：Tick 心跳、SleepTool 成本感知休眠、BriefTool 通信通道、外部消息唤醒、autoDream 记忆巩固。每个单独看不复杂，拼在一起构成完整的 AI 运行时——像极了 Linux 的 systemd。')

add_heading('SleepTool：成本感知的休眠', level=2)
add_image('01_tick-sleep成本权衡.png')
add_para('传统后台进程要么轮询要么事件驱动。KAIROS 两者都不是——它是成本驱动的。醒来一次花一次 API 调用的钱，但超过 5 分钟不醒来 prompt cache 过期，下次更贵。AI 需要自己算这笔账。')
add_para('有人实测 7x24 运行 AI Agent：切 Sonnet 做日常、Opus 做复杂推理，三个月月均 $187。SleepTool 本质上是把成本优化决策交给了 AI 自己。')

add_heading('BriefTool：思考和说话分离', level=2)
add_para('KAIROS 模式下，AI 所有用户可见输出必须通过 SendUserMessage（BriefTool）发送。普通文本输出对用户不可见。AI 在后台默默干活，只有觉得有信息值得你知道时才"说话"。区分 normal（回复提问）和 proactive（主动汇报）两种状态。')

add_heading('autoDream：AI 做梦', level=2)
add_image('03_autoDream四阶段.png')
add_para('autoDream 是后台记忆巩固进程。触发条件极其克制：24 小时 + 5 个新会话 + 锁文件互斥。以 fork 子进程运行，四阶段流程：定位已有记忆 -> 采集近期信号 -> 整合矛盾 -> 修剪索引（控制 200 行/25KB）。')
add_para('fork 子进程有严格工具限制：Bash 只读，Edit/Write 仅限 memory 目录。连记忆整理都不给 AI 完全自由——可控框架内的有限自主性。')

add_heading('趋同进化')
add_image('04_进化链.png')
add_para('多个独立开发者在不知道 KAIROS 存在的情况下做出了几乎一样的架构。约束条件决定了最优解：AI 要后台运行就需要心跳，需要心跳就要管理成本，管理成本就需要休眠，休眠需要唤醒条件，唤醒后要记住状态就需要记忆系统。鲨鱼和海豚长得像，不是谁抄了谁，是水的阻力决定了最优体型。')

add_heading('没人算过的账')
add_para('按 KAIROS 设计，5 分钟一次 tick，一天 288 次。即使每次 $0.05，一天 $14.4，一个月 $432。Devin 的前车之鉴：$500/月定价，15% 任务完成率，每个成功任务 $3,333。always-on agent 的定价模型还没有正确答案。')

add_heading('安全边界在哪里')
add_para('KAIROS 的心跳是自生成的，没有外部触发器可以检测。如果 Agent 被注入恶意指令，它会在下次 tick 时自动执行。企业安全团队需要回答：记忆文件存哪里？autoDream 生成的"事实"有没有人审核？这些不是代码问题，是治理问题。')

add_heading('对 AI 从业者意味着什么')
add_para('KAIROS 证明了 AI 编程助手的下一个战场不在模型能力，在运行模式。从 request-response 到 always-on，不是增加了功能，是改变了交互范式。做 AI Agent 产品的可以直接借鉴五层架构；评估工具的要看 always-on 进化潜力；企业安全负责人从现在开始准备 always-on agent 安全框架。')

# ============================================================
#  关键词
# ============================================================
doc.add_page_break()
add_heading('本期关键词')
add_keyword('KAIROS', '源码中 150+ 次引用的 feature flag，代表 AI 后台守护进程模式，通过心跳/休眠/记忆巩固让 AI 永不下线')
add_keyword('Feature Flag（特性标志）', '通过代码开关控制功能激活。Claude Code 用编译时死代码消除，被关闭的功能物理上从包里删除')
add_keyword('autoDream（自动做梦）', '记忆巩固子系统，四阶段流程。触发条件：24h + 5 sessions + 无并发')
add_keyword('Prompt Cache（提示缓存）', 'API 服务端缓存请求前缀，有效期约 5 分钟。KAIROS 成本模型的核心参数')
add_keyword('Daemon（守护进程）', '操作系统后台持续运行的程序。KAIROS 被类比为"AI 的 systemd"')
add_keyword('ACU（Agent Compute Unit）', 'Devin 的计费单位，约 15 分钟工作量。反映 always-on agent 定价难题')

# ============================================================
#  引用
# ============================================================
add_heading('引用')
refs = [
    '1. Claude Code v2.1.88 泄露源码（TypeScript）-- 一手素材',
    '2. VentureBeat -- Claude Code source code leak 报道',
    '3. TheNewStack -- swarms, daemons, and 44 feature flags',
    '4. kingy.ai -- KAIROS 安全性分析',
    '5. dev.to -- 7x24 AI Agent 运行成本实测（3个月 $560）',
    '6. understandingdata.com -- AI Daemons: Task vs Role 框架',
    '7. Reddit r/artificial -- 趋同进化讨论',
    '8. rumor -- 从Claude Code源码看Anthropic的产品野心',
]
for ref in refs:
    add_para(ref)

# ============================================================
#  保存
# ============================================================
out_name = '【AI笔记0407】KAIROS 7x24后台Daemon的野心'
out_path = os.path.join(OUT_DIR, f'{out_name}.docx')
doc.save(out_path)
print(f'OK: {out_path}')
